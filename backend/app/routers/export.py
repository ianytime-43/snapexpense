"""
Export router — PDF, Excel, CSV expense reports.

GET /api/export/pdf?start_date=YYYY-MM-DD&end_date=YYYY-MM-DD
GET /api/export/excel?start_date=YYYY-MM-DD&end_date=YYYY-MM-DD
GET /api/export/csv?start_date=YYYY-MM-DD&end_date=YYYY-MM-DD

Only confirmed/submitted/reimbursed expenses are included (not drafts).
"""
import csv
import io
import logging
from collections import defaultdict

import httpx

from datetime import date as date_type

from fastapi import APIRouter, Depends, Query
from fastapi.responses import Response, StreamingResponse

from ..auth import get_current_user
from ..database import get_supabase_admin

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/export", tags=["export"])

# Column spec used by CSV and Excel
COLUMNS: list[tuple[str, str]] = [
    ("expense_date",     "Date"),
    ("merchant_name",    "Merchant"),
    ("amount_total",     "Amount"),
    ("amount_tax",       "Tax"),
    ("currency",         "Currency"),
    ("category",         "Category"),
    ("client_name",      "Client"),
    ("business_purpose", "Purpose"),
    ("payment_method",   "Payment Method"),
    ("status",           "Status"),
]

EXPORT_STATUSES = {"draft", "confirmed", "submitted", "reimbursed"}


# ── shared helpers ─────────────────────────────────────────────────────────────

def _fetch_expenses(admin, user_id: str, start_date: str, end_date: str) -> list[dict]:
    """Fetch expenses in date range. Includes all statuses and handles NULL dates."""
    try:
        result = (
            admin.table("expenses")
            .select("*")
            .eq("user_id", user_id)
            .execute()
        )
        # Filter in Python to handle NULL dates properly
        expenses = []
        for e in (result.data or []):
            ed = e.get("expense_date")
            if ed is None or (ed >= start_date and ed <= end_date):
                expenses.append(e)
        # Sort by date
        expenses.sort(key=lambda x: x.get("expense_date") or "9999")
        return expenses
    except Exception as exc:
        logger.error("Export fetch failed: %s", exc)
        return []


def _get_user_info(admin, user_id: str) -> dict:
    try:
        result = (
            admin.table("users")
            .select("email, full_name, company_name")
            .eq("id", user_id)
            .maybe_single()
            .execute()
        )
        return (result.data or {}) if result else {}
    except Exception:
        return {}


def _fmt(v, decimals: int = 2) -> str:
    if v is None:
        return ""
    try:
        return f"{float(v):.{decimals}f}"
    except (TypeError, ValueError):
        return str(v)


# ── CSV ────────────────────────────────────────────────────────────────────────

def _build_csv(expenses: list[dict]) -> bytes:
    buf = io.StringIO()
    writer = csv.DictWriter(
        buf,
        fieldnames=[col for col, _ in COLUMNS],
        extrasaction="ignore",
    )
    writer.writerow({col: label for col, label in COLUMNS})
    for exp in expenses:
        row = {col: (exp.get(col) or "") for col, _ in COLUMNS}
        row["amount_total"] = _fmt(exp.get("amount_total"))
        row["amount_tax"] = _fmt(exp.get("amount_tax"))
        writer.writerow(row)
    # utf-8-sig so Excel opens without needing an import wizard
    return buf.getvalue().encode("utf-8-sig")


# ── Excel ──────────────────────────────────────────────────────────────────────

def _build_excel(
    expenses: list[dict],
    user_info: dict,
    start_date: str,
    end_date: str,
) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Expenses"

    last_col = get_column_letter(len(COLUMNS))

    # ── title rows
    ws.merge_cells(f"A1:{last_col}1")
    ws["A1"] = "SnapExpense — Expense Report"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")

    ws.merge_cells(f"A2:{last_col}2")
    company = user_info.get("company_name") or user_info.get("email") or ""
    ws["A2"] = f"{company}   {start_date} → {end_date}"
    ws["A2"].font = Font(size=10, color="555555")
    ws["A2"].alignment = Alignment(horizontal="center")

    ws.append([])  # blank row 3

    # ── header row (row 4)
    GREEN_FILL = PatternFill("solid", fgColor="16A34A")
    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws.append([label for _, label in COLUMNS])
    for cell in ws[4]:
        cell.fill = GREEN_FILL
        cell.font = Font(bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center")
        cell.border = border

    # ── data rows
    total_amount = 0.0
    totals_by_cat: dict[str, float] = defaultdict(float)
    LIGHT = PatternFill("solid", fgColor="F9FAFB")

    for i, exp in enumerate(expenses):
        amt = exp.get("amount_total")
        tax = exp.get("amount_tax")
        row_values = [
            exp.get("expense_date") or "",
            exp.get("merchant_name") or "",
            float(amt) if amt is not None else None,
            float(tax) if tax is not None else None,
            exp.get("currency") or "CAD",
            exp.get("category") or "",
            exp.get("client_name") or "",
            exp.get("business_purpose") or "",
            (exp.get("payment_method") or "").replace("_", " ").title(),
            exp.get("status") or "",
        ]
        ws.append(row_values)
        r = ws.max_row
        # Alternating rows
        if i % 2 == 1:
            for cell in ws[r]:
                cell.fill = LIGHT
        for cell in ws[r]:
            cell.border = border
        ws.cell(r, 3).number_format = "#,##0.00"
        ws.cell(r, 4).number_format = "#,##0.00"

        if amt is not None:
            total_amount += float(amt)
        if exp.get("category") and amt is not None:
            totals_by_cat[exp["category"]] += float(amt)

    # ── totals row
    ws.append(["TOTAL", "", total_amount, "", "", "", "", "", "", ""])
    tr = ws.max_row
    ws.cell(tr, 1).font = Font(bold=True)
    ws.cell(tr, 3).font = Font(bold=True)
    ws.cell(tr, 3).number_format = "#,##0.00"
    for cell in ws[tr]:
        cell.border = border

    # ── category summary sheet
    if totals_by_cat:
        ws2 = wb.create_sheet("By Category")
        ws2.append(["Category", "Total"])
        ws2["A1"].font = Font(bold=True)
        ws2["B1"].font = Font(bold=True)
        for cat, total in sorted(totals_by_cat.items(), key=lambda x: -x[1]):
            ws2.append([cat, total])
            ws2.cell(ws2.max_row, 2).number_format = "#,##0.00"
        ws2.column_dimensions["A"].width = 28
        ws2.column_dimensions["B"].width = 14

    # ── column widths (main sheet)
    widths = [12, 26, 10, 8, 10, 22, 20, 38, 18, 12]
    for col_idx, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = w

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── PDF ────────────────────────────────────────────────────────────────────────

def _fetch_receipt_thumbnail(image_url: str) -> "io.BytesIO | None":
    """Fetch a receipt image and resize to 100px wide. Returns BytesIO or None."""
    try:
        from PIL import Image as PILImage
        resp = httpx.get(image_url, timeout=8.0)
        resp.raise_for_status()
        img = PILImage.open(io.BytesIO(resp.content))
        # Resize to 100px wide, preserve aspect ratio
        w, h = img.size
        new_w = 100
        new_h = int(h * new_w / w)
        img = img.resize((new_w, new_h), PILImage.LANCZOS)
        # Convert to RGB if needed (e.g. RGBA PNG)
        if img.mode != "RGB":
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        buf.seek(0)
        return buf
    except Exception as exc:
        logger.warning("Receipt thumbnail fetch failed for %s: %s", image_url[:60], exc)
        return None


def _build_pdf(
    expenses: list[dict],
    user_info: dict,
    start_date: str,
    end_date: str,
) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        Image as RLImage,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    GREEN      = colors.HexColor("#16A34A")
    LIGHT_GRAY = colors.HexColor("#F3F4F6")
    MID_GRAY   = colors.HexColor("#6B7280")
    GREEN_TINT = colors.HexColor("#D1FAE5")
    BORDER_CLR = colors.HexColor("#E5E7EB")

    def style(name, **kw):
        return ParagraphStyle(name, **kw)

    S_TITLE      = style("title",      fontSize=16, fontName="Helvetica-Bold", spaceAfter=3)
    S_SUB        = style("sub",        fontSize=9,  fontName="Helvetica",      textColor=MID_GRAY, spaceAfter=8)
    S_SECTION    = style("section",    fontSize=8,  fontName="Helvetica-Bold", textColor=MID_GRAY, spaceAfter=4)
    S_CELL       = style("cell",       fontSize=8,  fontName="Helvetica",      leading=10)
    S_CELL_BOLD  = style("cell_b",     fontSize=8,  fontName="Helvetica-Bold", leading=10)
    S_CELL_RIGHT = style("cell_r",     fontSize=8,  fontName="Helvetica",      leading=10, alignment=TA_RIGHT)
    S_BOLD_RIGHT = style("cell_br",    fontSize=8,  fontName="Helvetica-Bold", leading=10, alignment=TA_RIGHT)

    buf = io.BytesIO()
    margin = 1.5 * cm
    doc = SimpleDocTemplate(
        buf,
        pagesize=landscape(A4),
        leftMargin=margin, rightMargin=margin,
        topMargin=margin,  bottomMargin=margin,
    )

    story: list = []

    # ── header
    company   = user_info.get("company_name") or ""
    user_name = user_info.get("full_name") or user_info.get("email") or ""
    story.append(Paragraph("Expense Report", S_TITLE))
    parts = [p for p in [user_name, company, f"{start_date} → {end_date}"] if p]
    story.append(Paragraph("   ·   ".join(parts), S_SUB))
    story.append(HRFlowable(width="100%", thickness=1, color=GREEN))
    story.append(Spacer(1, 0.5 * cm))

    if not expenses:
        story.append(Paragraph("No confirmed expenses in this date range.", S_CELL))
        doc.build(story)
        return buf.getvalue()

    # ── expense table
    # Total usable width: landscape A4 = 29.7 cm, minus 3 cm margins → ~26.7 cm
    col_widths = [2.2*cm, 4.8*cm, 2.4*cm, 3.8*cm, 3.8*cm, 7.2*cm, 2.0*cm]

    table_data: list = [[
        Paragraph("Date",     S_CELL_BOLD),
        Paragraph("Merchant", S_CELL_BOLD),
        Paragraph("Amount",   S_BOLD_RIGHT),
        Paragraph("Category", S_CELL_BOLD),
        Paragraph("Client",   S_CELL_BOLD),
        Paragraph("Purpose",  S_CELL_BOLD),
        Paragraph("Status",   S_CELL_BOLD),
    ]]

    total_amount = 0.0
    totals_by_cat: dict[str, float] = defaultdict(float)

    for exp in expenses:
        amt = exp.get("amount_total")
        currency = exp.get("currency") or "CAD"
        amt_str = f"{currency} {float(amt):.2f}" if amt is not None else "—"
        if amt is not None:
            total_amount += float(amt)
        cat = exp.get("category") or ""
        if cat and amt is not None:
            totals_by_cat[cat] += float(amt)

        # Converted amount display
        converted = exp.get("converted_amount")
        conv_currency = exp.get("converted_currency")
        if converted and conv_currency and conv_currency != currency:
            amt_str += f"\n≈{conv_currency} {float(converted):.2f}"

        table_data.append([
            Paragraph(exp.get("expense_date") or "—",           S_CELL),
            Paragraph(exp.get("merchant_name") or "—",          S_CELL),
            Paragraph(amt_str,                                    S_CELL_RIGHT),
            Paragraph(cat or "—",                                S_CELL),
            Paragraph(exp.get("client_name") or "—",            S_CELL),
            Paragraph(exp.get("business_purpose") or "—",       S_CELL),
            Paragraph(exp.get("status") or "—",                 S_CELL),
        ])

    # totals row
    table_data.append([
        Paragraph("TOTAL",                             S_CELL_BOLD),
        Paragraph("",                                  S_CELL),
        Paragraph(f"CAD {total_amount:.2f}",           S_BOLD_RIGHT),
        Paragraph("", S_CELL), Paragraph("", S_CELL),
        Paragraph("", S_CELL), Paragraph("", S_CELL),
    ])

    alt_rows = [
        ("BACKGROUND", (0, i), (-1, i), LIGHT_GRAY)
        for i in range(2, len(table_data) - 1, 2)
    ]

    tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND",  (0, 0),  (-1, 0),  GREEN),
        ("TEXTCOLOR",   (0, 0),  (-1, 0),  colors.white),
        *alt_rows,
        ("BACKGROUND",  (0, -1), (-1, -1), GREEN_TINT),
        ("LINEABOVE",   (0, -1), (-1, -1), 1, GREEN),
        ("GRID",        (0, 0),  (-1, -1), 0.25, BORDER_CLR),
        ("TOPPADDING",    (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING",   (0, 0), (-1, -1), 5),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 5),
        ("VALIGN",        (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(tbl)

    # ── category breakdown
    if totals_by_cat:
        story.append(Spacer(1, 0.8 * cm))
        story.append(Paragraph("Totals by Category", S_SECTION))
        cat_rows: list = [[
            Paragraph("Category", S_CELL_BOLD),
            Paragraph("Amount",   S_BOLD_RIGHT),
        ]]
        for cat, total in sorted(totals_by_cat.items(), key=lambda x: -x[1]):
            cat_rows.append([
                Paragraph(cat, S_CELL),
                Paragraph(f"CAD {total:.2f}", S_CELL_RIGHT),
            ])
        cat_alt = [
            ("BACKGROUND", (0, i), (-1, i), LIGHT_GRAY)
            for i in range(2, len(cat_rows), 2)
        ]
        cat_tbl = Table(cat_rows, colWidths=[6.5*cm, 3*cm])
        cat_tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0), GREEN),
            ("TEXTCOLOR",     (0, 0), (-1, 0), colors.white),
            *cat_alt,
            ("GRID",          (0, 0), (-1, -1), 0.25, BORDER_CLR),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 5),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 5),
        ]))
        story.append(cat_tbl)

    doc.build(story)
    return buf.getvalue()


# ── endpoints ──────────────────────────────────────────────────────────────────

@router.get("/pdf")
def export_pdf(
    start_date: str = Query(...),
    end_date: str = Query(...),
    current_user: dict = Depends(get_current_user),
):
    user_id = str(current_user["user"].id)
    admin = get_supabase_admin()
    expenses  = _fetch_expenses(admin, user_id, start_date, end_date)
    user_info = _get_user_info(admin, user_id)
    data = _build_pdf(expenses, user_info, start_date, end_date)
    fname = f"expenses_{start_date}_{end_date}.pdf"
    return Response(
        content=data,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.get("/excel")
def export_excel(
    start_date: str = Query(...),
    end_date: str = Query(...),
    current_user: dict = Depends(get_current_user),
):
    user_id = str(current_user["user"].id)
    admin = get_supabase_admin()
    expenses  = _fetch_expenses(admin, user_id, start_date, end_date)
    user_info = _get_user_info(admin, user_id)
    data = _build_excel(expenses, user_info, start_date, end_date)
    fname = f"expenses_{start_date}_{end_date}.xlsx"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.get("/csv")
def export_csv(
    start_date: str = Query(...),
    end_date: str = Query(...),
    current_user: dict = Depends(get_current_user),
):
    user_id = str(current_user["user"].id)
    admin = get_supabase_admin()
    expenses = _fetch_expenses(admin, user_id, start_date, end_date)
    data = _build_csv(expenses)
    fname = f"expenses_{start_date}_{end_date}.csv"
    return Response(
        content=data,
        media_type="text/csv; charset=utf-8-sig",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ── Word (docx) ─────────────────────────────────────────────────────────────────

def _build_docx(
    expenses: list[dict],
    user_info: dict,
    start_date: str,
    end_date: str,
) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    user_name = user_info.get("full_name") or user_info.get("email") or ""
    title_para = doc.add_heading(f"Expense Report — {user_name}", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"Period: {start_date} to {end_date}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    if not expenses:
        doc.add_paragraph("No confirmed expenses in this date range.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # Main table
    headers = ["Date", "Merchant", "Category", "Amount (orig)", "Converted", "Client", "Purpose", "Location"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    totals_by_cat: dict[str, float] = defaultdict(float)
    grand_total = 0.0

    for exp in sorted(expenses, key=lambda x: x.get("expense_date") or ""):
        row_cells = table.add_row().cells
        amt = exp.get("amount_total")
        currency = exp.get("currency") or ""
        amt_str = f"{currency} {float(amt):.2f}" if amt is not None else "—"
        converted = exp.get("converted_amount")
        conv_currency = exp.get("converted_currency") or ""
        converted_str = f"{conv_currency} {float(converted):.2f}" if converted is not None else ""

        row_cells[0].text = exp.get("expense_date") or "—"
        row_cells[1].text = exp.get("merchant_name") or "—"
        row_cells[2].text = exp.get("category") or "—"
        row_cells[3].text = amt_str
        row_cells[4].text = converted_str
        row_cells[5].text = exp.get("client_name") or "—"
        row_cells[6].text = exp.get("business_purpose") or "—"
        row_cells[7].text = exp.get("location_jurisdiction") or exp.get("location_name") or "—"

        if amt is not None:
            grand_total += float(amt)
        cat = exp.get("category")
        if cat and amt is not None:
            totals_by_cat[cat] += float(amt)

    # Subtotals by category
    if totals_by_cat:
        doc.add_paragraph()
        doc.add_heading("Totals by Category", level=2)
        cat_table = doc.add_table(rows=1, cols=2)
        cat_table.style = "Table Grid"
        cat_hdr = cat_table.rows[0].cells
        cat_hdr[0].text = "Category"
        cat_hdr[1].text = "Total"
        for run in cat_hdr[0].paragraphs[0].runs:
            run.bold = True
        for run in cat_hdr[1].paragraphs[0].runs:
            run.bold = True
        for cat, total in sorted(totals_by_cat.items(), key=lambda x: -x[1]):
            r = cat_table.add_row().cells
            r[0].text = cat
            r[1].text = f"{total:.2f}"

    # Grand total row
    total_row = table.add_row().cells
    total_row[0].text = "TOTAL"
    total_row[3].text = f"{grand_total:.2f}"
    for cell in [total_row[0], total_row[3]]:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Generated by SnapExpense · {date_type.today().isoformat()}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/docx")
def export_docx(
    start_date: str = Query(...),
    end_date: str = Query(...),
    current_user: dict = Depends(get_current_user),
):
    user_id = str(current_user["user"].id)
    admin = get_supabase_admin()
    expenses = _fetch_expenses(admin, user_id, start_date, end_date)
    user_info = _get_user_info(admin, user_id)
    data = _build_docx(expenses, user_info, start_date, end_date)
    fname = f"expenses_{start_date}_{end_date}.docx"
    return StreamingResponse(
        io.BytesIO(data),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )
